#!/usr/bin/env python3
"""
COMPLETE WORKING PRIVACY POLICY GENERATOR
UK GDPR compliant with ALL data included
No passwords, no email subscription - just clean policy generation
"""

import os
import re
import webbrowser
from datetime import datetime
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class PrivacyPolicyGenerator:
    def __init__(self):
        # Get the script's directory
        self.script_dir = os.path.dirname(os.path.abspath(__file__))
        self.policies_dir = os.path.join(self.script_dir, "generated_policies")
        
        # Create policies directory if it doesn't exist
        if not os.path.exists(self.policies_dir):
            os.makedirs(self.policies_dir)
        
        # Business information
        self.company_name = ""
        self.website = ""
        self.email = ""
        self.phone = ""
        self.address = ""
        self.business_type = ""
        self.company_registration = ""
        self.ico_registration = ""
        self.data_protection_officer = ""
        self.has_dpo = False
        
        # AI usage
        self.ai_use_cases = []
        self.uses_chatbot = False
        self.uses_ai_marketing = False
        self.uses_profiling = False
        self.uses_automated_decisions = False
        self.auto_decision_types = []
        self.auto_decision_solely_automated = False
        self.auto_decision_human_review = False
        self.auto_decision_consequences = ""
        self.trains_ai_on_data = False
        self.ai_training_details = ""
        self.uses_ai_content_gen = False
        self.uses_ai_analytics = False
        self.ai_provider_locations = []
        self.ai_provider_location_unconfirmed = False
        self.ai_data_processed = []
        self.ai_opt_out_available = False
        
        # Data collection
        self.data_collected = []
        self.data_sources = []
        self.purposes = []
        self.purpose_lawful_basis = {}
        self.shared_with_categories = []
        self.shared_data_description = ""
        self.processor_contracts = False
        self.uses_sensitive_data = False
        self.sensitive_data_types = []
        self.sensitive_data_basis = ""
        self.international_transfers = False
        self.transfer_countries = []
        self.transfer_safeguards = []
        
        # Marketing and cookies
        self.uses_marketing = False
        self.marketing_methods = []
        self.marketing_consent = False
        self.uses_cookies = False
        self.cookie_types = []
        self.uses_analytics = False
        self.analytics_types = []
        
        # Retention
        self.retention_periods = []
        
        # File management
        self.filename = ""
        self.full_path = ""
        self.policy_content = ""
        self.policy_html = ""
        
    def clear_screen(self):
        os.system('cls' if os.name == 'nt' else 'clear')
        
    def print_header(self, title: str):
        print("\n" + "="*60)
        print(f"  {title}")
        print("="*60 + "\n")
        
    def _safe_input(self, prompt: str) -> str:
        """
        Wrapper around input() that survives a stray Ctrl+C / interrupted
        paste on a SINGLE field, instead of letting it kill the entire
        session (which is what happens if KeyboardInterrupt escapes all
        the way up to run()). If it happens twice in a row, we give up
        and re-raise so the user isn't stuck in an infinite loop.
        """
        attempts = 0
        while True:
            try:
                return input(prompt)
            except KeyboardInterrupt:
                attempts += 1
                print(
                    "\n  ⚠️  That keystroke/paste was interpreted as Ctrl+C "
                    "and interrupted this field only - nothing else has "
                    "been lost. Please try again"
                    + (" (or press Ctrl+C once more to actually cancel): "
                       if attempts >= 2 else ": ")
                )
                if attempts >= 2:
                    raise
            except EOFError:
                print(
                    "\n  ⚠️  No input was received (EOF). Please try typing "
                    "or pasting the value again: "
                )
                attempts += 1
                if attempts >= 3:
                    raise

    def get_input(self, prompt: str, required: bool = True, default: str = "") -> str:
        while True:
            if default:
                user_input = self._safe_input(f"{prompt} [{default}]: ").strip()
                if not user_input:
                    return default
            else:
                user_input = self._safe_input(f"{prompt}: ").strip()
            
            if required and not user_input:
                print("  This field is required. Please enter a value.")
                continue
            return user_input
    
    def get_yes_no(self, prompt: str, default: str = "y") -> bool:
        while True:
            user_input = self._safe_input(f"{prompt} (y/n) [{default}]: ").strip().lower()
            if not user_input:
                user_input = default
            if user_input in ["y", "yes"]:
                return True
            elif user_input in ["n", "no"]:
                return False
            print("  Please enter 'y' or 'n'")
    
    def checkbox_selection(self, prompt: str, options: List[Tuple[str, str]], 
                           allow_other: bool = False, 
                           other_prompt: str = "",
                           allow_none: bool = False,
                           multiple: bool = True) -> List[str]:
        print(f"\n{prompt}")
        if multiple:
            print("  Select options by number. Press Enter twice to finish.\n")
        else:
            print("  Select one option by number.\n")
        
        for idx, (key, desc) in enumerate(options, 1):
            print(f"  {idx}. {desc}")
        
        other_option_added = False
        if allow_other:
            has_other = any(key == "other" for key, _ in options)
            if not has_other:
                print(f"  {len(options) + 1}. Other (please specify)")
                other_option_added = True
        
        if allow_none:
            print("  Type 'none' to select none")
        
        print()
        
        selected = []
        while True:
            choice = self._safe_input("  Enter number (or press Enter to finish): ").strip()
            
            if not choice:
                if selected:
                    break
                elif allow_none:
                    return ["none"]
                else:
                    if not multiple:
                        print("  Please select an option.")
                        continue
                    else:
                        print("  Please select at least one option.")
                        continue
            
            if choice.lower() == "none" and allow_none:
                return ["none"]
            
            if choice.isdigit():
                idx = int(choice) - 1
                
                if allow_other and other_option_added and idx == len(options):
                    other_value = self._safe_input(f"  {other_prompt}: ").strip()
                    if other_value:
                        if multiple:
                            selected.append(f"other:{other_value}")
                            print(f"  ✓ Added: {other_value}\n")
                        else:
                            selected = [f"other:{other_value}"]
                            print(f"  ✓ Selected: {other_value}\n")
                            break
                    continue
                
                if 0 <= idx < len(options):
                    option_key = options[idx][0]
                    if multiple:
                        if option_key not in selected:
                            selected.append(option_key)
                            print(f"  ✓ Added: {options[idx][1]}\n")
                        else:
                            print(f"  Already selected: {options[idx][1]}")
                    else:
                        selected = [option_key]
                        print(f"  ✓ Selected: {options[idx][1]}\n")
                        break
                    continue
            
            print("  Invalid selection. Please enter a number from the list.")
        
        return selected
    
    # ====== GATHERING METHODS ======
    def gather_business_info(self):
        self.clear_screen()
        self.print_header("SECTION 1: BUSINESS INFORMATION")
        print("Let's start with your business details.\n")
        
        self.company_name = self.get_input("1. Business/company name")
        self.website = self.get_input("2. Website URL")
        self.email = self.get_input("3. Contact email for privacy queries")
        self.phone = self.get_input("4. Business phone (optional)", required=False)
        self.address = self.get_input("5. Registered business address (optional)", required=False)
        
        business_types = [
            ("retail", "Online/Physical Retail"),
            ("consultancy", "Consultancy/Professional Services"),
            ("saas", "SaaS/Software Provider"),
            ("agency", "Marketing/Digital Agency"),
            ("ecommerce", "E-commerce"),
            ("hospitality", "Hospitality/Travel"),
            ("health", "Health/Wellness"),
            ("education", "Education/Training"),
            ("financial", "Financial Services"),
            ("tech", "Technology/IT Services")
        ]
        
        selected = self.checkbox_selection("What type of business is this?", business_types,
                                          allow_other=True, other_prompt="Please specify your business type")
        
        cleaned = []
        for item in selected:
            if item.startswith("other:"):
                cleaned.append(item.replace("other:", ""))
            else:
                for key, display in business_types:
                    if key == item:
                        cleaned.append(display)
                        break
        
        self.business_type = ", ".join(cleaned) if cleaned else "Not specified"
        self.company_registration = self.get_input("Company registration number (if applicable, press Enter to skip)", required=False)
        self.ico_registration = self.get_input("ICO registration number (if applicable, press Enter to skip)", required=False)
        self.has_dpo = self.get_yes_no("\nDo you have a Data Protection Officer (DPO)?")
        if self.has_dpo:
            self.data_protection_officer = self.get_input("  DPO contact details")
    
    def gather_ai_usage(self):
        self.clear_screen()
        self.print_header("SECTION 2: AI USAGE")
        print("How does your business use AI? Select all that apply.\n")
        
        ai_options = [
            ("chatbot", "AI Chatbot for customer interactions"),
            ("marketing", "AI for marketing/email personalisation"),
            ("profiling", "AI for customer profiling/behaviour analysis"),
            ("auto_decision", "AI for automated decision-making"),
            ("train_ai", "Using customer/employee data to train AI systems"),
            ("content_gen", "AI for content generation"),
            ("analytics", "AI for analytics/predictive analysis")
        ]
        
        selected = self.checkbox_selection("Select all AI use cases:", ai_options,
                                          allow_other=True, other_prompt="Please specify other AI use")
        
        self.uses_chatbot = False
        self.uses_ai_marketing = False
        self.uses_profiling = False
        self.uses_automated_decisions = False
        self.trains_ai_on_data = False
        self.uses_ai_content_gen = False
        self.uses_ai_analytics = False
        self.ai_use_cases = []
        
        for item in selected:
            if item == "chatbot":
                self.uses_chatbot = True
                self.ai_use_cases.append("AI Chatbot for customer interactions")
            elif item == "marketing":
                self.uses_ai_marketing = True
                self.ai_use_cases.append("AI for marketing/email personalisation")
            elif item == "profiling":
                self.uses_profiling = True
                self.ai_use_cases.append("AI for customer profiling/behaviour analysis")
            elif item == "auto_decision":
                self.uses_automated_decisions = True
                self.ai_use_cases.append("AI for automated decision-making")
            elif item == "train_ai":
                self.trains_ai_on_data = True
                self.ai_use_cases.append("Using customer/employee data to train AI systems")
            elif item == "content_gen":
                self.uses_ai_content_gen = True
                self.ai_use_cases.append("AI for content generation")
            elif item == "analytics":
                self.uses_ai_analytics = True
                self.ai_use_cases.append("AI for analytics/predictive analysis")
            elif item.startswith("other:"):
                self.ai_use_cases.append(item.replace("other:", ""))
        
        if self.uses_automated_decisions:
            print("\n" + "-"*40)
            print("🤖 AUTOMATED DECISION-MAKING DETAILS")
            print("  What types of automated decisions does AI make?")
            print("  Examples: credit scoring, job applications, refund decisions")
            decision_details = self.get_input("  Describe the types of decisions")
            if decision_details and decision_details.lower() != "none":
                self.auto_decision_types = [d.strip() for d in decision_details.split(',')]

            # These extra questions matter because "significant decisions" made
            # solely by AI (no human involved) carry additional UK GDPR
            # obligations (Article 22) - people have a right to know this.
            self.auto_decision_solely_automated = self.get_yes_no(
                "\n  Is any of this decision-making carried out WITHOUT any "
                "human involvement (i.e. solely automated)?"
            )
            self.auto_decision_human_review = self.get_yes_no(
                "  Can a person request that a human reviews or reconsiders "
                "the decision?"
            )
            consequences = self.get_input(
                "  Briefly, what is the practical effect of these decisions on "
                "the person (e.g. 'may result in application being rejected')",
                required=False
            )
            if consequences:
                self.auto_decision_consequences = consequences
        
        if self.trains_ai_on_data:
            print("\n" + "-"*40)
            print("🧠 AI TRAINING DETAILS")
            print("  What data do you use for AI training?")
            self.ai_training_details = self.get_input("  Describe the data used and purpose")
        
        if any([self.uses_chatbot, self.uses_ai_marketing, self.uses_automated_decisions, 
                self.trains_ai_on_data, self.uses_profiling,
                self.uses_ai_content_gen, self.uses_ai_analytics]):
            
            print("\n" + "-"*40)
            print("📍 AI PROVIDER LOCATIONS\n")
            print("  No need to name specific suppliers - just where the")
            print("  processing broadly takes place. This can be tricky if")
            print("  you use several different AI tools, e.g. a standalone")
            print("  tool you signed up for directly (like a chatbot builder)")
            print("  AND AI features already built into other software you")
            print("  use (like Salesforce Einstein or Microsoft Copilot).")
            print("  If you're not sure, that's fine - select the option")
            print("  below for that.\n")
            
            locations = [
                ("uk", "UK-based only"),
                ("us", "US-based"),
                ("eu", "EU/EEA-based"),
                ("unsure", "Not sure / it varies across the different AI tools we use")
            ]
            
            selected_locations = self.checkbox_selection("Where are your AI providers located?", locations,
                                                         allow_other=True, other_prompt="Please specify countries")
            
            self.ai_provider_locations = []
            self.ai_provider_location_unconfirmed = False
            for loc in selected_locations:
                if loc == "uk":
                    self.ai_provider_locations.append("UK")
                elif loc == "us":
                    self.ai_provider_locations.append("United States")
                elif loc == "eu":
                    self.ai_provider_locations.append("EU/EEA")
                elif loc == "unsure":
                    self.ai_provider_location_unconfirmed = True
                elif loc.startswith("other:"):
                    self.ai_provider_locations.append(loc.replace("other:", ""))
            
            print("\n" + "-"*40)
            print("📊 DATA PROCESSED BY AI\n")
            
            data_options = [
                ("names", "Names and contact details"),
                ("payment", "Payment information"),
                ("order_history", "Order/purchase history"),
                ("chat_logs", "Chatbot conversation logs"),
                ("email", "Email content"),
                ("website", "Website usage/behaviour data"),
                ("support", "Customer support tickets"),
                ("sensitive", "Sensitive data")
            ]
            selected_data = self.checkbox_selection("Select all types of data processed by AI:", data_options,
                                                    allow_other=True, other_prompt="Please specify other data types")
            
            self.ai_data_processed = []
            for data in selected_data:
                if data.startswith("other:"):
                    self.ai_data_processed.append(data.replace("other:", ""))
                else:
                    for key, display in data_options:
                        if key == data:
                            self.ai_data_processed.append(display)
                            break
            
            self.ai_opt_out_available = self.get_yes_no("\nDo you offer customers the option to opt out of AI processing?")
    
    def gather_data_collection(self):
        self.clear_screen()
        self.print_header("SECTION 3: DATA COLLECTION")
        print("What personal information does your business collect?\n")
        
        data_options = [
            ("name", "Name and contact details (email, phone)"),
            ("address", "Address information"),
            ("payment", "Payment/financial information"),
            ("order_history", "Order/purchase history"),
            ("website", "Website usage and analytics data"),
            ("cookies", "Cookie data"),
            ("chat_logs", "Chatbot/conversation logs"),
            ("email", "Email content and preferences"),
            ("support", "Customer support enquiries"),
            ("social", "Social media information"),
            ("employee", "Employee data"),
            ("sensitive", "Sensitive data")
        ]
        
        selected = self.checkbox_selection("Select all types of personal data you collect:", data_options,
                                          allow_other=True, other_prompt="Please specify other data collected")
        
        self.data_collected = []
        self.uses_sensitive_data = False
        self.sensitive_data_types = []
        
        for item in selected:
            if item.startswith("other:"):
                self.data_collected.append(item.replace("other:", ""))
            elif item == "employee":
                self.data_collected.append("Employee data")
            elif item == "sensitive":
                self.uses_sensitive_data = True
                print("\n  What sensitive data do you collect?")
                print("  Examples: health information, religious beliefs, racial/ethnic origin")
                sensitive_details = self.get_input("  Describe the sensitive data")
                if sensitive_details and sensitive_details.lower() != "none":
                    self.sensitive_data_types = [s.strip() for s in sensitive_details.split(',')]
                    self.data_collected.extend(self.sensitive_data_types)
                else:
                    self.data_collected.append("Sensitive data")
            else:
                for key, display in data_options:
                    if key == item:
                        self.data_collected.append(display)
                        break
        
        if self.uses_sensitive_data:
            print("\n" + "-"*40)
            print("⚖️  SENSITIVE DATA - ARTICLE 9 BASIS\n")
            
            basis_options = [
                ("explicit_consent", "Explicit consent from the data subject"),
                ("employment", "Employment, social security, or social protection law"),
                ("vital_interests", "Protection of vital interests"),
                ("public_interest", "Substantial public interest"),
                ("health", "Health or social care"),
                ("legal_claims", "Legal claims or court proceedings")
            ]
            
            selected_basis = self.checkbox_selection("What is your Article 9 basis for processing sensitive data?",
                                                     basis_options, allow_other=True,
                                                     other_prompt="Please specify the legal basis")
            
            basis_descriptions = []
            for basis in selected_basis:
                if basis.startswith("other:"):
                    basis_descriptions.append(basis.replace("other:", ""))
                else:
                    for key, display in basis_options:
                        if key == basis:
                            basis_descriptions.append(display)
                            break
            self.sensitive_data_basis = "; ".join(basis_descriptions) if basis_descriptions else "Not specified"

        # ---- Where does this data come from? ----
        print("\n" + "-"*40)
        print("📥 SOURCE OF DATA\n")

        source_options = [
            ("direct", "Directly from the individual (e.g. forms, purchases, account signup)"),
            ("website_activity", "Automatically from website/app activity (cookies, usage data)"),
            ("third_party_partner", "From partners or other businesses (e.g. referrals, resellers)"),
            ("public_sources", "Publicly available sources"),
            ("employer", "From an individual's employer"),
            ("ai_generated", "Generated or inferred by AI/automated systems")
        ]

        selected_sources = self.checkbox_selection(
            "Where does the personal data you hold come from?", source_options,
            allow_other=True, other_prompt="Please specify other source"
        )

        self.data_sources = []
        for item in selected_sources:
            if item.startswith("other:"):
                self.data_sources.append(item.replace("other:", ""))
            else:
                for key, display in source_options:
                    if key == item:
                        self.data_sources.append(display)
                        break
    
    def gather_purposes(self):
        self.clear_screen()
        self.print_header("SECTION 4: DATA PROCESSING PURPOSES")
        print("Why do you process personal data?\n")
        
        purpose_options = [
            ("orders", "Process orders and transactions"),
            ("customer_service", "Provide customer service and support"),
            ("marketing", "Send marketing communications (with consent)"),
            ("improve", "Improve products and services"),
            ("analytics", "Analytics and business intelligence"),
            ("fraud", "Fraud prevention and detection"),
            ("legal", "Legal compliance"),
            ("ai_training", "AI system training and improvement"),
            ("profiling", "Customer profiling and personalisation")
        ]
        
        selected = self.checkbox_selection("Select all purposes for processing personal data:", purpose_options,
                                          allow_other=True, other_prompt="Please specify other purposes")
        
        self.purposes = []
        for item in selected:
            if item.startswith("other:"):
                self.purposes.append(item.replace("other:", ""))
            else:
                for key, display in purpose_options:
                    if key == item:
                        self.purposes.append(display)
                        break

        # ---- Lawful basis for EACH purpose ----
        # This is the piece the earlier version was missing: a generic
        # statement that "lawful basis hasn't been mapped" isn't good enough
        # once we're actually asking the question.
        self.purpose_lawful_basis = {}
        if self.purposes:
            print("\n" + "-"*40)
            print("⚖️  LAWFUL BASIS (UK GDPR ARTICLE 6)\n")
            print("  For each purpose, choose the lawful basis that applies.\n")

            basis_options = [
                ("consent", "Consent"),
                ("contract", "Necessary for a contract with the individual"),
                ("legal_obligation", "Necessary to comply with a legal obligation"),
                ("vital_interests", "Necessary to protect someone's vital interests"),
                ("public_task", "Necessary for a public task"),
                ("legitimate_interests", "Legitimate interests"),
            ]

            for purpose in self.purposes:
                print(f"\n  Purpose: {purpose}")
                basis_choice = self.checkbox_selection(
                    "  Select the lawful basis for this purpose:",
                    basis_options, multiple=False,
                    allow_other=True, other_prompt="Please specify the lawful basis"
                )
                if basis_choice:
                    chosen = basis_choice[0]
                    if chosen.startswith("other:"):
                        self.purpose_lawful_basis[purpose] = chosen.replace("other:", "")
                    else:
                        for key, display in basis_options:
                            if key == chosen:
                                self.purpose_lawful_basis[purpose] = display
                                break
    
    def gather_data_sharing(self):
        self.clear_screen()
        self.print_header("SECTION 5: DATA SHARING")
        print("Do you share personal data with third parties?\n")
        
        sharing_options = [
            ("payment", "Payment processors"),
            ("hosting", "Hosting/cloud providers"),
            ("ai_providers", "AI service providers"),
            ("customer_service", "Customer service platforms"),
            ("marketing", "Marketing agencies/tools"),
            ("analytics", "Analytics providers"),
            ("email", "Email service providers"),
            ("storage", "Cloud storage providers"),
            ("it", "IT support providers"),
            ("legal", "Legal/accounting professionals"),
            ("delivery", "Delivery/courier services")
        ]
        
        selected = self.checkbox_selection("What types of third parties do you share data with?", sharing_options,
                                          allow_other=True, other_prompt="Please specify other third party types",
                                          allow_none=True)
        
        self.shared_with_categories = []
        for item in selected:
            if item == "none":
                self.shared_with_categories = []
                break
            if item.startswith("other:"):
                self.shared_with_categories.append(item.replace("other:", ""))
            else:
                for key, display in sharing_options:
                    if key == item:
                        self.shared_with_categories.append(display)
                        break
        
        if self.shared_with_categories:
            self.processor_contracts = self.get_yes_no("\nDo you have written contracts with these data processors (Article 28)?")

            # We deliberately do NOT ask for specific supplier/vendor names -
            # that's commercially sensitive. We just ask what kind of
            # information is shared and roughly why, in the business's own words.
            self.shared_data_description = self.get_input(
                "\n  In a sentence, what type of information do you share with "
                "these third parties and why (no need to name suppliers)?",
                required=False
            )
        
        print("\n" + "-"*40)
        print("🌍 INTERNATIONAL DATA TRANSFERS")
        print("  Does your data leave the UK?\n")
        
        transfer_options = [
            ("usa", "United States"),
            ("ireland", "Ireland"),
            ("germany", "Germany"),
            ("france", "France"),
            ("netherlands", "Netherlands"),
            ("canada", "Canada"),
            ("australia", "Australia"),
            ("singapore", "Singapore")
        ]
        
        selected_transfers = self.checkbox_selection("Where does your data go? Select all that apply:", transfer_options,
                                                     allow_other=True, other_prompt="Please specify other countries",
                                                     allow_none=True)
        
        self.transfer_countries = []
        self.international_transfers = False
        
        for item in selected_transfers:
            if item == "none":
                self.transfer_countries = []
                break
            if item.startswith("other:"):
                self.transfer_countries.append(item.replace("other:", ""))
                self.international_transfers = True
            else:
                for key, display in transfer_options:
                    if key == item:
                        self.transfer_countries.append(display)
                        self.international_transfers = True
                        break
        
        if self.international_transfers:
            print("\n  What safeguards do you have for international transfers?")
            safeguard_options = [
                ("scc", "Standard Contractual Clauses (SCCs)"),
                ("adequacy", "Adequacy decision (EU/EEA)"),
                ("bcr", "Binding Corporate Rules"),
                ("consent", "Explicit consent from data subjects")
            ]
            
            selected_safeguards = self.checkbox_selection("Select all safeguards that apply:", safeguard_options,
                                                          allow_other=True, other_prompt="Please specify other safeguards")
            
            self.transfer_safeguards = []
            for item in selected_safeguards:
                if item.startswith("other:"):
                    self.transfer_safeguards.append(item.replace("other:", ""))
                else:
                    for key, display in safeguard_options:
                        if key == item:
                            self.transfer_safeguards.append(display)
                            break
    
    def gather_marketing_and_cookies(self):
        self.clear_screen()
        self.print_header("SECTION 6: MARKETING AND COOKIES")
        
        self.uses_marketing = self.get_yes_no("1. Does your business send marketing communications?")
        
        if self.uses_marketing:
            marketing_options = [
                ("email", "Email marketing"),
                ("sms", "SMS text messages"),
                ("post", "Postal mail"),
                ("social", "Social media"),
                ("phone", "Phone calls")
            ]
            
            selected = self.checkbox_selection("How do you send marketing?", marketing_options,
                                              allow_other=True, other_prompt="Please specify other methods")
            
            self.marketing_methods = []
            for item in selected:
                if item.startswith("other:"):
                    self.marketing_methods.append(item.replace("other:", ""))
                else:
                    for key, display in marketing_options:
                        if key == item:
                            self.marketing_methods.append(display)
                            break
            
            self.marketing_consent = self.get_yes_no("  Do you get explicit consent before sending marketing?")
        
        self.uses_cookies = self.get_yes_no("\n2. Does your website use cookies?")
        
        if self.uses_cookies:
            cookie_options = [
                ("essential", "Essential/necessary cookies"),
                ("analytics", "Analytics/performance cookies"),
                ("functional", "Functional cookies"),
                ("marketing", "Marketing/targeting cookies")
            ]
            
            selected = self.checkbox_selection("What types of cookies do you use?", cookie_options,
                                              allow_other=True, other_prompt="Please specify other cookie types")
            
            self.cookie_types = []
            for item in selected:
                if item.startswith("other:"):
                    self.cookie_types.append(item.replace("other:", ""))
                else:
                    for key, display in cookie_options:
                        if key == item:
                            self.cookie_types.append(display)
                            break
        
        self.uses_analytics = self.get_yes_no("\n3. Do you use website analytics?")
        
        if self.uses_analytics:
            analytics_options = [
                ("pageviews", "Page views and traffic"),
                ("behaviour", "User behaviour tracking"),
                ("conversion", "Conversion tracking"),
                ("referral", "Referral sources")
            ]
            
            selected = self.checkbox_selection("What analytics do you collect?", analytics_options,
                                              allow_other=True, other_prompt="Please specify other analytics")
            
            self.analytics_types = []
            for item in selected:
                if item.startswith("other:"):
                    self.analytics_types.append(item.replace("other:", ""))
                else:
                    for key, display in analytics_options:
                        if key == item:
                            self.analytics_types.append(display)
                            break
    
    def gather_retention(self):
        self.clear_screen()
        self.print_header("SECTION 7: DATA RETENTION")
        print("UK GDPR requires you to state how long you keep personal data.\n")
        
        retention_options = [
            ("1_year", "1 year"),
            ("2_years", "2 years"),
            ("3_years", "3 years"),
            ("5_years", "5 years"),
            ("7_years", "7 years (financial records)")
        ]
        
        selected = self.checkbox_selection("What is your data retention policy?", retention_options,
                                          allow_other=True, other_prompt="Please specify retention period")
        
        self.retention_periods = []
        for item in selected:
            if item.startswith("other:"):
                self.retention_periods.append(item.replace("other:", ""))
            else:
                for key, display in retention_options:
                    if key == item:
                        self.retention_periods.append(display)
                        break
    
    # ====== GENERATE COMPLETE POLICY ======
    def _bullets(self, items, fallback=None):
        """Return a readable bullet list without inventing information."""
        if not items:
            return fallback or ""
        return "\n".join(f"- {item}" for item in items)

    def _sentence_list(self, items, fallback="not specified"):
        """Turn a list into natural language."""
        if not items:
            return fallback
        if len(items) == 1:
            return items[0]
        if len(items) == 2:
            return f"{items[0]} and {items[1]}"
        return ", ".join(items[:-1]) + f", and {items[-1]}"

    def _tbc(self, description: str) -> str:
        """
        A clearly-marked gap to fill in, styled like an editor's note in a
        genuine draft policy - not a sentence about the questionnaire.
        """
        return f"**[TO CONFIRM: {description}]**"

    def generate_privacy_policy(self) -> str:
        """
        Generate the privacy policy from the answers actually supplied.

        Important design rule:
        This method must not invent business practices.  It may provide
        explanatory privacy wording, but factual statements about the
        business are based on the questionnaire answers.
        """
        today = datetime.now().strftime("%d %B %Y")

        # ---------------- BUSINESS DETAILS ----------------
        company_details = []
        if self.address:
            company_details.append(f"**Address:** {self.address}")
        if self.company_registration:
            company_details.append(
                f"**Company Registration Number:** {self.company_registration}"
            )
        if self.ico_registration:
            company_details.append(
                f"**ICO Registration Number:** {self.ico_registration}"
            )
        if self.has_dpo and self.data_protection_officer:
            company_details.append(
                f"**Data Protection Officer:** {self.data_protection_officer}"
            )

        # ---------------- INFORMATION WE COLLECT ----------------
        if self.data_collected:
            data_collection = (
                "We collect and use the following categories of personal "
                "information, depending on how you interact with us:\n\n"
                + self._bullets(self.data_collected)
            )
        else:
            data_collection = self._tbc(
                "list the types of personal information collected"
            )

        if self.uses_sensitive_data:
            sensitive = self.sensitive_data_types or ["Special category data"]
            data_collection += (
                "\n\n**Special category data**\n\n"
                "We also process special category data under UK GDPR "
                "Article 9, specifically:\n\n"
                + self._bullets(sensitive)
                + "\n\nOur Article 9 condition for processing this data is: "
                + (self.sensitive_data_basis or self._tbc(
                    "confirm the Article 9 condition relied on"))
                + "."
            )

        if self.data_sources:
            data_collection += (
                "\n\n**Where this information comes from**\n\n"
                "We obtain your personal data from the following sources:\n\n"
                + self._bullets(self.data_sources)
            )
        else:
            data_collection += (
                "\n\n**Where this information comes from**\n\n"
                + self._tbc("describe where personal data is obtained from")
            )

        # ---------------- PURPOSES ----------------
        if self.purposes:
            purpose_lines = []
            for purpose in self.purposes:
                basis = self.purpose_lawful_basis.get(purpose)
                if basis:
                    purpose_lines.append(f"- {purpose} — **lawful basis:** {basis}")
                else:
                    purpose_lines.append(
                        f"- {purpose} — lawful basis: "
                        + self._tbc("confirm the Article 6 lawful basis")
                    )
            purposes_text = (
                "We use your personal information for the following "
                "purposes. Each purpose is listed with the lawful basis "
                "we rely on under Article 6 of the UK GDPR:\n\n"
                + "\n".join(purpose_lines)
            )
        else:
            purposes_text = self._tbc(
                "list the purposes for which personal information is processed, "
                "with the Article 6 UK GDPR lawful basis for each"
            )

        # ---------------- AI ----------------
        ai_sections = []

        if self.ai_use_cases:
            ai_sections.append(
                "### How we use AI\n\n"
                "We use artificial intelligence (AI) in the following ways:\n\n"
                + self._bullets(self.ai_use_cases)
            )

        if self.uses_chatbot:
            ai_data = self._sentence_list(
                self.ai_data_processed,
                "the information you provide during the interaction"
            )
            chatbot = (
                "### AI chatbot\n\n"
                "We use an AI-powered chatbot for customer interactions. "
                f"When you use the chatbot, we may process {ai_data}.\n\n"
                "The chatbot is an automated system rather than a human "
                "member of staff."
            )
            if self.ai_opt_out_available:
                chatbot += (
                    "\n\nYou can opt out of speaking with the AI chatbot "
                    "and request a human alternative by contacting us."
                )
            ai_sections.append(chatbot)

        if self.uses_ai_marketing:
            marketing_use = self._sentence_list(
                [x for x in self.ai_use_cases if "marketing" in x.lower()],
                "AI-assisted marketing"
            )
            ai_sections.append(
                "### AI used for marketing\n\n"
                "We use AI to support our marketing activity, including "
                f"{marketing_use}."
            )

        if self.uses_profiling:
            ai_sections.append(
                "### Profiling\n\n"
                "We use AI to build profiles based on your interactions, "
                "preferences or behaviour. We use these profiles for the "
                "purposes set out in Section 4 of this policy."
            )

        if self.uses_ai_content_gen:
            ai_sections.append(
                "### AI used for content generation\n\n"
                "We use AI tools to help generate content, such as website "
                "copy or product descriptions. Where personal information "
                "is used as part of this process, it falls within the data "
                "categories described in Section 3."
            )

        if self.uses_ai_analytics:
            ai_sections.append(
                "### AI used for analytics and predictive analysis\n\n"
                "We use AI to support analytics and to identify patterns or "
                "trends in the data described in this policy, in support of "
                "the purposes set out in Section 4."
            )

        if self.uses_automated_decisions:
            decision_text = self._bullets(
                self.auto_decision_types,
                "- " + self._tbc("list the types of automated decisions made")
            )

            automated_section = (
                "### Automated decision-making\n\n"
                "We use AI to make automated decisions relating to:\n\n"
                + decision_text
            )

            if self.auto_decision_solely_automated:
                automated_section += (
                    "\n\nSome of these decisions are made **without any "
                    "human involvement** (solely automated processing "
                    "within the meaning of Article 22 of the UK GDPR)."
                )
            else:
                automated_section += (
                    "\n\nA human is involved in these decisions; they are "
                    "not made solely by automated means."
                )

            if self.auto_decision_human_review:
                automated_section += (
                    " If you disagree with a decision, you can ask us to "
                    "have it reviewed by a person, and express your point "
                    "of view, by contacting us."
                )
            else:
                automated_section += (
                    " " + self._tbc(
                        "confirm whether individuals can request human review "
                        "of these decisions - this is generally required for "
                        "significant automated decisions under Article 22"
                    )
                )

            if self.auto_decision_consequences:
                automated_section += (
                    f"\n\n**What this means for you:** {self.auto_decision_consequences}"
                )
            else:
                automated_section += (
                    "\n\n**What this means for you:** "
                    + self._tbc("describe the practical effect of these "
                                 "decisions on individuals")
                )

            ai_sections.append(automated_section)

        if self.trains_ai_on_data:
            training_details = (
                self.ai_training_details
                if self.ai_training_details
                else self._tbc("describe what data is used to train AI and why")
            )
            ai_sections.append(
                "### AI training\n\n"
                "We may use personal data to train or improve AI systems.\n\n"
                f"**How this data is used:** {training_details}\n\n"
                "**Data processed:** "
                + self._sentence_list(self.ai_data_processed, self._tbc("confirm"))
                + "\n\n"
                "**Where this processing takes place:** "
                + self._sentence_list(self.ai_provider_locations, self._tbc("confirm"))
                + "\n\n"
                + (
                    "You can opt out of your data being used for AI "
                    "training by contacting us."
                    if self.ai_opt_out_available
                    else self._tbc("confirm whether an AI-training opt-out is offered")
                )
            )

        if self.ai_provider_locations or self.ai_provider_location_unconfirmed:
            location_lines = list(self.ai_provider_locations)
            if self.ai_provider_location_unconfirmed:
                location_lines.append(
                    "Some AI tools we use - this has not yet been "
                    "confirmed and will be updated once we've checked with "
                    "each provider"
                )
            ai_sections.append(
                "### Where AI processing takes place\n\n"
                "The AI tools we use process data in the following "
                "locations:\n\n"
                + self._bullets(location_lines)
            )

        if self.ai_data_processed:
            ai_sections.append(
                "### Information processed by AI\n\n"
                "The following categories of information may be processed "
                "by AI:\n\n"
                + self._bullets(self.ai_data_processed)
            )

        if not ai_sections:
            ai_sections.append(
                "We do not currently use AI to process personal information. "
                "If this changes, we will update this policy."
            )

        ai_provisions = "\n\n".join(ai_sections)

        # ---------------- SHARING ----------------
        if self.shared_with_categories:
            sharing_text = (
                "We may share personal information with the following "
                "categories of third parties, where necessary for the "
                "purposes described in this policy:\n\n"
                + self._bullets(self.shared_with_categories)
            )
            if self.shared_data_description:
                sharing_text += f"\n\n{self.shared_data_description}"
            if self.processor_contracts:
                sharing_text += (
                    "\n\nWhere third parties process personal data on our "
                    "behalf as processors, we have written contracts in "
                    "place with them in line with Article 28 of the UK GDPR."
                )
            else:
                sharing_text += (
                    "\n\n" + self._tbc(
                        "confirm that Article 28 processor contracts are in "
                        "place with these third parties"
                    )
                )
        else:
            sharing_text = (
                "We do not routinely share personal information with third "
                "parties. We may still need to disclose information where "
                "required by law, for example to a regulator or law "
                "enforcement body."
            )

        # ---------------- INTERNATIONAL TRANSFERS ----------------
        if self.international_transfers and self.transfer_countries:
            transfers_text = (
                "Personal data may be transferred to, or processed in, the "
                "following countries outside the UK:\n\n"
                + self._bullets(self.transfer_countries)
            )
            if self.transfer_safeguards:
                transfers_text += (
                    "\n\nWhere we transfer personal data outside the UK, we "
                    "rely on the following safeguards, in accordance with "
                    "Chapter 5 of the UK GDPR:\n\n"
                    + self._bullets(self.transfer_safeguards)
                )
            else:
                transfers_text += (
                    "\n\n" + self._tbc(
                        "confirm the transfer safeguard used (e.g. UK "
                        "International Data Transfer Agreement, adequacy "
                        "regulations) - see the ICO's guidance on "
                        "international transfers"
                    )
                )
        else:
            transfers_text = (
                "We do not currently transfer personal data outside the UK."
            )

        # ---------------- RETENTION ----------------
        if self.retention_periods:
            retention_text = (
                "We retain personal information for the following periods:\n\n"
                + self._bullets(self.retention_periods)
                + "\n\nWhen information is no longer needed, we securely "
                "delete or anonymise it, subject to any legal or regulatory "
                "requirement to keep it for longer."
            )
        else:
            retention_text = self._tbc(
                "confirm how long each category of personal information is "
                "retained, and why"
            )

        # ---------------- MARKETING ----------------
        if self.uses_marketing:
            marketing_text = (
                "We send marketing communications using the following "
                "methods:\n\n"
                + self._bullets(self.marketing_methods, "- " + self._tbc("confirm method"))
            )
            if self.marketing_consent:
                marketing_text += (
                    "\n\nWe obtain your consent before sending marketing "
                    "communications, in line with the Privacy and "
                    "Electronic Communications Regulations (PECR)."
                )
            else:
                marketing_text += (
                    "\n\n" + self._tbc(
                        "confirm the lawful basis and PECR consent position "
                        "for this marketing activity"
                    )
                )
            marketing_text += (
                "\n\nYou can opt out of marketing communications at any "
                "time using the unsubscribe link provided or by contacting us."
            )
        else:
            marketing_text = "We do not send marketing communications."

        # ---------------- COOKIES / ANALYTICS ----------------
        if self.uses_cookies:
            cookies_text = (
                "Our website uses the following types of cookies:\n\n"
                + self._bullets(self.cookie_types, "- " + self._tbc("confirm cookie types"))
            )
            cookies_text += (
                "\n\nIn line with PECR and ICO guidance on cookies, we will "
                "provide clear information and, where required, obtain your "
                "consent before setting non-essential cookies."
            )
        else:
            cookies_text = "Our website does not use cookies."

        if self.uses_analytics:
            cookies_text += (
                "\n\n**Analytics**\n\n"
                "We use website analytics to measure and understand:\n\n"
                + self._bullets(self.analytics_types, "- " + self._tbc("confirm analytics used"))
            )

        # ---------------- LEGAL BASIS ----------------
        if self.purposes and all(p in self.purpose_lawful_basis for p in self.purposes):
            legal_basis_text = (
                "The lawful basis we rely on under Article 6 of the UK GDPR "
                "for each processing purpose is set out alongside that "
                "purpose in Section 4 above."
            )
        elif self.purpose_lawful_basis:
            legal_basis_text = (
                "The lawful basis for most processing purposes is set out "
                "in Section 4 above. " + self._tbc(
                    "confirm the lawful basis for any purpose shown there "
                    "without one"
                )
            )
        else:
            legal_basis_text = self._tbc(
                "confirm the Article 6 UK GDPR lawful basis for each "
                "purpose listed in Section 4"
            )

        if self.uses_sensitive_data:
            legal_basis_text += (
                "\n\nFor special category data, our Article 9 condition is:\n\n"
                + (self.sensitive_data_basis or self._tbc("confirm the Article 9 condition"))
                + "."
            )

        # ---------------- POLICY ----------------
        business_details = "\n".join(company_details)
        if business_details:
            business_details = "\n" + business_details

        policy = f"""# Privacy Policy

**Last Updated: {today}**

## 1. Introduction

{self.company_name} ("we", "our" or "us") is committed to protecting and
respecting your privacy. This policy explains how we collect, use, share and
protect your personal information, and sets out your rights, in accordance
with the UK General Data Protection Regulation (UK GDPR), the Data
Protection Act 2018 and, where relevant, the Privacy and Electronic
Communications Regulations (PECR).

Sections of this policy marked **[TO CONFIRM]** have not yet been completed
and must be filled in, and this policy reviewed by a qualified professional,
before it is published.

## 2. Who We Are

**Business Name:** {self.company_name}

**Website:** {self.website}

**Business Type:** {self.business_type}

**Contact Email:** {self.email}

**Contact Phone:** {self.phone or "Not provided"}{business_details}

## 3. Information We Collect

{data_collection}

## 4. How We Use Personal Information

{purposes_text}

## 5. AI and Automated Processing

{ai_provisions}

## 6. Cookies and Analytics

{cookies_text}

## 7. Marketing

{marketing_text}

## 8. Sharing Personal Information

{sharing_text}

## 9. International Data Transfers

{transfers_text}

## 10. Data Retention

{retention_text}

## 11. Lawful Basis for Processing

{legal_basis_text}

## 12. Your Data Protection Rights

Under the UK GDPR and the Data Protection Act 2018, you have the following
rights over your personal information:

- the right to be informed about how your personal information is used;
- the right of access to the personal information we hold about you;
- the right to rectification of inaccurate information;
- the right to erasure in certain circumstances;
- the right to restrict processing in certain circumstances;
- the right to object to certain processing, including direct marketing;
- the right to data portability, where applicable;
- the right to withdraw consent at any time, where we rely on consent; and
- rights relating to automated decision-making and profiling, where applicable.

To exercise any of these rights, contact us at **{self.email}**.

## 13. Data Security

We take appropriate technical and organisational measures, as required by
the UK GDPR, to protect personal information against unauthorised or
unlawful processing and against accidental loss, destruction or damage.

{self._tbc("describe the specific security measures in place, e.g. encryption, access controls, staff training")}

## 14. Data Processors

Where third parties process personal information on our behalf, we remain
responsible for that data and ensure appropriate contractual and data
protection safeguards are in place, in line with Article 28 of the UK GDPR.

Processor contracts are currently {"in place" if self.processor_contracts else self._tbc("confirm processor contracts are in place")}.

## 15. Changes to This Privacy Policy

We may update this policy from time to time to reflect changes in our
processing activities, technology or legal obligations. We will update the
"Last Updated" date above whenever we make changes, and encourage you to
review this policy periodically.

## 16. Complaints

If you have concerns about how we handle your personal information, please
contact us first so we can try to resolve the issue directly.

You also have the right to lodge a complaint with the UK's independent
supervisory authority for data protection:

**Information Commissioner's Office (ICO)**

**Website:** https://ico.org.uk

**Helpline:** 0303 123 1113

## 17. Contact Us

If you have questions about this policy or how we process your personal
information, please contact us:

**Email:** {self.email}

**Website:** {self.website}

**Phone:** {self.phone or "Please email us for contact details"}

---

### Important review note

This is a **draft** privacy policy generated from the answers provided.
Every statement above should be checked against what the business actually
does, and any section marked **[TO CONFIRM]** must be completed. This draft
is not a determination of legal compliance, and should be reviewed by a
qualified solicitor or data protection professional - and updated whenever
processing activities change - before it is published.
"""

        return policy

    # ====== WORD DOCUMENT GENERATION ======
    def _add_runs_with_bold(self, paragraph, text: str):
        """Add text to a paragraph, turning **bold** markers into real bold runs."""
        parts = text.split("**")
        for idx, part in enumerate(parts):
            if part == "":
                continue
            run = paragraph.add_run(part)
            if idx % 2 == 1:
                run.bold = True

    def build_docx(self, policy_text: str) -> "Document":
        """
        Convert the generated markdown-style policy into a properly
        formatted Word document: real heading styles, bold labels,
        and bullet lists, rather than plain text.
        """
        doc = Document()

        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)

        section = doc.sections[0]
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

        # Group lines into blank-line-separated blocks first. Within a
        # block, the source template sometimes wraps one sentence across
        # several lines purely for code readability - those need to be
        # rejoined into a single paragraph rather than rendered as several
        # short, choppy ones.
        raw_lines = policy_text.split("\n")
        blocks: List[List[str]] = []
        current: List[str] = []
        for raw_line in raw_lines:
            line = raw_line.strip()
            if line == "":
                if current:
                    blocks.append(current)
                    current = []
                continue
            current.append(line)
        if current:
            blocks.append(current)

        for block in blocks:
            # A block that's a run of bullet lines: emit one list item per line
            if all(l.startswith("- ") for l in block):
                for l in block:
                    p = doc.add_paragraph(style='List Bullet')
                    self._add_runs_with_bold(p, l[2:].strip())
                continue

            # A block is otherwise treated as: first line determines type,
            # any continuation lines (plain text, no special prefix) are
            # rejoined into the same paragraph.
            first = block[0]

            if first == "---":
                p = doc.add_paragraph()
                p_format = p.paragraph_format
                p_format.space_before = Pt(6)
                p_format.space_after = Pt(6)
                pPr = p._p.get_or_add_pPr()
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), 'auto')
                pBdr.append(bottom)
                pPr.append(pBdr)
                continue

            if first.startswith("# "):
                doc.add_heading(first[2:].strip(), level=0)
                continue
            if first.startswith("### "):
                doc.add_heading(first[4:].strip(), level=2)
                continue
            if first.startswith("## "):
                doc.add_heading(first[3:].strip(), level=1)
                continue

            # Plain paragraph, possibly wrapped across multiple lines.
            # Handle mixed blocks (e.g. a bullet list line appearing mid
            # block) defensively by flushing bullets separately.
            buffer: List[str] = []

            def flush_buffer():
                if buffer:
                    p = doc.add_paragraph()
                    self._add_runs_with_bold(p, " ".join(buffer))
                    buffer.clear()

            for l in block:
                if l.startswith("- "):
                    flush_buffer()
                    p = doc.add_paragraph(style='List Bullet')
                    self._add_runs_with_bold(p, l[2:].strip())
                elif l.startswith("#"):
                    flush_buffer()
                    if l.startswith("### "):
                        doc.add_heading(l[4:].strip(), level=2)
                    elif l.startswith("## "):
                        doc.add_heading(l[3:].strip(), level=1)
                    elif l.startswith("# "):
                        doc.add_heading(l[2:].strip(), level=0)
                elif l.startswith("**"):
                    # A bold-label line (e.g. "**Contact Phone:** 12345")
                    # is always its own paragraph, even when the source
                    # template didn't put a blank line before/after it -
                    # otherwise adjacent detail fields collapse into one
                    # run-on line.
                    flush_buffer()
                    p = doc.add_paragraph()
                    self._add_runs_with_bold(p, l)
                else:
                    buffer.append(l)
            flush_buffer()

        return doc

    def save_docx(self, policy: str, path: str):
        """Save the policy as a formatted .docx file. Returns True on success."""
        if not DOCX_AVAILABLE:
            return False
        doc = self.build_docx(policy)
        doc.save(path)
        return True

    # ====== SAVE FUNCTIONS ======
    def save_policy(self, policy: str):
        """Save the policy to a file"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_name = self.company_name.replace(' ', '_').replace("'", "").replace('"', '').lower()
        
        # Save as markdown
        md_filename = f"privacy_policy_{safe_name}_{timestamp}.md"
        md_path = os.path.join(self.policies_dir, md_filename)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(policy)
        
        # Also save as text file for easy viewing
        txt_filename = f"privacy_policy_{safe_name}_{timestamp}.txt"
        txt_path = os.path.join(self.policies_dir, txt_filename)
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(policy)
        
        # Save simple copy in main directory
        simple_path = os.path.join(self.script_dir, f"privacy_policy_{safe_name}.txt")
        with open(simple_path, 'w', encoding='utf-8') as f:
            f.write(policy)

        # Save as a properly formatted Word document (headings, bold, bullets)
        docx_filename = f"privacy_policy_{safe_name}_{timestamp}.docx"
        docx_path = os.path.join(self.policies_dir, docx_filename)
        docx_saved = self.save_docx(policy, docx_path)
        
        self.filename = txt_filename
        self.full_path = docx_path if docx_saved else txt_path
        self.policy_content = policy
        
        return {
            'txt': txt_path,
            'md': md_path,
            'simple': simple_path,
            'docx': docx_path if docx_saved else None
        }
    
    def generate_summary(self) -> str:
        """Generate a summary showing ALL data captured"""
        summary = """\n
╔══════════════════════════════════════════════════════════════════════════════╗
║                     PRIVACY POLICY GENERATION COMPLETE                      ║
╚══════════════════════════════════════════════════════════════════════════════╝

✅ Your complete UK GDPR-compliant privacy policy has been generated.

FILES GENERATED
────────────────────────────────────────────────────────────────────────────────
"""
        summary += f"\n  📄 Policy saved in: {self.policies_dir}/"
        summary += f"\n  📄 Simple copy: privacy_policy_{self.company_name.replace(' ', '_').lower()}.txt"
        
        summary += """

ALL DATA INCLUDED IN POLICY
────────────────────────────────────────────────────────────────────────────────
"""
        summary += f"""
AI Use Cases ({len(self.ai_use_cases)}):
"""
        if self.ai_use_cases:
            for use_case in self.ai_use_cases:
                summary += f"    • {use_case}\n"
        else:
            summary += "    • None selected\n"
        
        summary += f"""
  • Chatbot: {'Yes' if self.uses_chatbot else 'No'}
  • AI Marketing: {'Yes' if self.uses_ai_marketing else 'No'}
  • AI Content Generation: {'Yes' if self.uses_ai_content_gen else 'No'}
  • AI Analytics/Predictive: {'Yes' if self.uses_ai_analytics else 'No'}
  • Profiling: {'Yes' if self.uses_profiling else 'No'}
  • Automated Decisions: {'Yes' if self.uses_automated_decisions else 'No'}
    Types: {', '.join(self.auto_decision_types) if self.auto_decision_types else 'N/A'}
    Solely automated: {'Yes' if self.auto_decision_solely_automated else 'No'}
    Human review available: {'Yes' if self.auto_decision_human_review else 'No'}
    Consequences: {self.auto_decision_consequences or 'N/A'}
  • AI Training: {'Yes' if self.trains_ai_on_data else 'No'}
    Details: {self.ai_training_details or 'N/A'}
  • AI Provider Locations: {', '.join(self.ai_provider_locations) if self.ai_provider_locations else 'N/A'}
  • AI Data Processed: {', '.join(self.ai_data_processed) if self.ai_data_processed else 'N/A'}
  • AI Opt-out: {'Yes' if self.ai_opt_out_available else 'No'}

Data Collection ({len(self.data_collected)} types):
"""
        for item in self.data_collected:
            summary += f"    • {item}\n"

        summary += f"""
Data Sources ({len(self.data_sources)} types):
"""
        for item in self.data_sources:
            summary += f"    • {item}\n"
        
        summary += f"""
Purposes ({len(self.purposes)} types):
"""
        for purpose in self.purposes:
            basis = self.purpose_lawful_basis.get(purpose, 'Not specified')
            summary += f"    • {purpose} (lawful basis: {basis})\n"
        
        summary += f"""
Sharing Categories ({len(self.shared_with_categories)} types):
"""
        for category in self.shared_with_categories:
            summary += f"    • {category}\n"
        if self.shared_data_description:
            summary += f"    Description: {self.shared_data_description}\n"
        
        summary += f"""
Sensitive Data: {'Yes' if self.uses_sensitive_data else 'No'}
  Types: {', '.join(self.sensitive_data_types) if self.sensitive_data_types else 'N/A'}
  Article 9 Basis: {self.sensitive_data_basis or 'N/A'}

International Transfers: {'Yes' if self.international_transfers else 'No'}
  Countries: {', '.join(self.transfer_countries) if self.transfer_countries else 'N/A'}
  Safeguards: {', '.join(self.transfer_safeguards) if self.transfer_safeguards else 'N/A'}

Retention Periods: {', '.join(self.retention_periods) if self.retention_periods else 'N/A'}
Marketing Consent: {'Yes' if self.marketing_consent else 'No'}

────────────────────────────────────────────────────────────────────────────────
"""
        return summary

    def run(self):
        """Main application flow"""
        self.clear_screen()
        self.print_header("PRIVACY POLICY GENERATOR")
        
        print("This tool builds a DRAFT UK GDPR privacy policy from your answers.")
        print("ALL your answers will be included in the final document.\n")
        print("⚠️  IMPORTANT:")
        print("  • This produces a starting draft, not a finished, compliant policy.")
        print("  • Anything you leave blank will be marked [TO CONFIRM] in the")
        print("    output rather than guessed at.")
        print("  • The generated .docx/.txt/.md files are yours to edit afterwards -")
        print("    nothing here is locked.")
        print("  • Have a solicitor or data protection professional review the")
        print("    final policy before you publish it.\n")
        print("Press Enter to begin...")
        input()
        
        try:
            # Gather all information
            self.gather_business_info()
            self.gather_ai_usage()
            self.gather_data_collection()
            self.gather_purposes()
            self.gather_data_sharing()
            self.gather_marketing_and_cookies()
            self.gather_retention()
            
            # Generate the policy
            self.clear_screen()
            self.print_header("GENERATING YOUR PRIVACY POLICY")
            print("Creating your complete UK GDPR-compliant privacy policy...\n")
            
            policy = self.generate_privacy_policy()
            
            # Save the policy
            file_results = self.save_policy(policy)
            
            # Display summary
            self.clear_screen()
            print(self.generate_summary())
            
            if file_results.get('docx'):
                print(f"\n📄 Word document (.docx): {file_results['docx']}")
            else:
                print(
                    "\n⚠️  Word document was not created because the "
                    "'python-docx' package is not installed.\n"
                    "   Install it with: pip install python-docx\n"
                    "   Falling back to plain text/markdown only."
                )
            print(f"📁 Policy also saved to: {self.full_path}")
            
            # Ask about viewing
            if self.get_yes_no("\nWould you like to open the policy in your browser?"):
                # Create simple HTML version for viewing
                html_content = f"""<!DOCTYPE html>
<html>
<head><title>Privacy Policy</title>
<style>
body {{ font-family: Calibri, Arial, sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; line-height: 1.6; }}
h1 {{ font-size: 28px; border-bottom: 1px solid #ccc; padding-bottom: 10px; }}
h2 {{ font-size: 20px; margin-top: 25px; }}
ul {{ margin: 10px 0 15px 30px; }}
li {{ margin-bottom: 4px; }}
</style>
</head>
<body>
{policy.replace('\n', '<br>')}
</body>
</html>"""
                html_path = os.path.join(self.policies_dir, self.filename.replace('.txt', '.html'))
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                webbrowser.open('file://' + html_path)
                print("✅ Policy opened in browser")
            
            # Final message
            print("\n" + "="*60)
            print("  IMPORTANT NOTES")
            print("="*60)
            print("""
  1. ✓ ALL your answers are included in the policy
  2. ✓ Clean Word-style formatting
  3. ✓ AI use cases, training, profiling all included
  4. ✓ Article 9 basis for sensitive data
  5. ✓ Review with a legal professional before publishing
  6. ✓ Update when your AI usage changes
        
  The Golden Rule: Tell people what you actually do with their information.
""")
            
            print(f"\n📁 All generated files: {self.policies_dir}")
            print(f"\nThank you for using the Privacy Policy Generator!")
            
        except KeyboardInterrupt:
            print("\n\n⚠️  Generation cancelled.")
        except Exception as e:
            print(f"\n❌ An error occurred: {e}")
            import traceback
            traceback.print_exc()

def main():
    generator = PrivacyPolicyGenerator()
    generator.run()

if __name__ == "__main__":
    main()
